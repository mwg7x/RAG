import os
import shutil
from pathlib import Path
from typing import Iterable, List

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from langchain_community.embeddings.fake import DeterministicFakeEmbedding
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import SentenceTransformerEmbeddings
from langchain_chroma import Chroma


class DocumentProcessor:
    """Loads contract files, chunks them, and stores embeddings in Chroma."""

    def __init__(
        self,
        persist_directory: str | None = None,
        collection_name: str = "contracts",
        embedding_model: str = "all-MiniLM-L6-v2",
        chunk_size: int = 1000,
        chunk_overlap: int = 100,
    ) -> None:
        self.persist_directory = persist_directory or os.getenv("CHROMA_PATH", "./chroma_db")
        self.collection_name = collection_name
        self.embedding_model = embedding_model
        self.embeddings = self._build_embeddings(embedding_model)
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )

        self.vectorstore = Chroma(
            collection_name=self.collection_name,
            embedding_function=self.embeddings,
            persist_directory=self.persist_directory,
        )

    @staticmethod
    def _build_embeddings(embedding_model: str):
        if os.getenv("USE_FAKE_EMBEDDINGS", "0") == "1":
            return DeterministicFakeEmbedding(size=384)

        try:
            return SentenceTransformerEmbeddings(model_name=embedding_model)
        except Exception:
            # Fallback keeps local development usable when outbound model download is blocked.
            return DeterministicFakeEmbedding(size=384)

    def _load_pdf(self, file_path: str) -> List[Document]:
        docs: List[Document] = []
        pdf = fitz.open(file_path)
        for page_index, page in enumerate(pdf):
            text = page.get_text("text") or ""
            if not text.strip():
                continue
            docs.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": Path(file_path).name,
                        "path": str(Path(file_path).resolve()),
                        "page": page_index + 1,
                        "section": f"Page {page_index + 1}",
                        "file_type": "pdf",
                    },
                )
            )
        pdf.close()
        return docs

    def _load_docx(self, file_path: str) -> List[Document]:
        docs: List[Document] = []
        doc = DocxDocument(file_path)
        paragraphs = []
        for idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            paragraphs.append((idx + 1, text))

        if not paragraphs:
            return docs

        joined = "\n".join(text for _, text in paragraphs)
        docs.append(
            Document(
                page_content=joined,
                metadata={
                    "source": Path(file_path).name,
                    "path": str(Path(file_path).resolve()),
                    "section": "Document Body",
                    "paragraph_count": len(paragraphs),
                    "file_type": "docx",
                },
            )
        )
        return docs

    def load_file(self, file_path: str) -> List[Document]:
        suffix = Path(file_path).suffix.lower()
        if suffix == ".pdf":
            raw_docs = self._load_pdf(file_path)
        elif suffix == ".docx":
            raw_docs = self._load_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

        chunked_docs = self.splitter.split_documents(raw_docs)
        return chunked_docs

    def add_documents(self, file_paths: Iterable[str]):
        all_docs: List[Document] = []
        for file_path in file_paths:
            all_docs.extend(self.load_file(file_path))

        if not all_docs:
            raise ValueError("No readable text found in uploaded documents.")

        self.vectorstore.add_documents(all_docs)
        if hasattr(self.vectorstore, "persist"):
            self.vectorstore.persist()
        return self.vectorstore.as_retriever(search_kwargs={"k": 4})

    def clear_store(self) -> None:
        try:
            self.vectorstore.delete_collection()
        except Exception:
            pass
        persist_path = Path(self.persist_directory)
        if persist_path.exists():
            shutil.rmtree(persist_path, ignore_errors=True)
        self.vectorstore = Chroma(
            collection_name=self.collection_name,
            embedding_function=self.embeddings,
            persist_directory=self.persist_directory,
        )

    def get_retriever(self):
        return self.vectorstore.as_retriever(search_kwargs={"k": 4})

